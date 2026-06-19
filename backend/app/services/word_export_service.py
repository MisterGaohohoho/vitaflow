from pathlib import Path
from typing import Any

from docx import Document
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.export_record import ExportRecord
from app.models.resume import Resume


def _plain(value: Any) -> str:
    if isinstance(value, list):
        lines = []
        for item in value:
            if isinstance(item, dict):
                text = " / ".join(_inline(part) for part in item.values() if _inline(part))
            else:
                text = str(item).lstrip("- ").strip()
            if text:
                lines.append(text)
        return "\n".join(lines)
    if isinstance(value, dict):
        return "\n".join(f"{key}：{_inline(item)}" for key, item in value.items() if _inline(item))
    return str(value or "").replace("**", "").replace("#", "").strip()


def _inline(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return " / ".join(_inline(item) for item in value if _inline(item))
    if isinstance(value, dict):
        return " / ".join(_inline(item) for item in value.values() if _inline(item))
    return str(value).replace("**", "").replace("#", "").strip()


def export_word(db: Session, user_id: int, resume: Resume) -> Path:
    settings.export_path.mkdir(parents=True, exist_ok=True)
    data = resume.resume_data or {}
    basics = data.get("basics", {})
    doc = Document()
    doc.add_heading(_inline(basics.get("name")) or resume.title, 0)
    doc.add_paragraph(" | ".join(filter(None, [_inline(basics.get("title")), _inline(basics.get("phone")), _inline(basics.get("email")), _inline(basics.get("location"))])))

    titles = data.get("layout", {}).get("section_titles", {})
    custom_sections = {
        item["id"]: item
        for item in data.get("custom_sections", [])
        if isinstance(item, dict) and item.get("id")
    }
    for key in data.get("layout", {}).get("section_order", []):
        if key == "basics" or key in data.get("layout", {}).get("hidden_sections", []):
            continue
        section = custom_sections.get(key) if key in custom_sections else data.get(key)
        if not section:
            continue
        doc.add_heading(titles.get(key, key), level=1)
        if key == "summary":
            doc.add_paragraph(_plain(section.get("content")))
        elif isinstance(section, list):
            for item in section:
                doc.add_heading(_inline(item.get("name") or item.get("company") or item.get("school") or item.get("title")) or "经历", level=2)
                meta = " / ".join(
                    filter(
                        None,
                        [
                            _inline(item.get("role")),
                            _inline(item.get("position")),
                            _inline(item.get("major")),
                            _inline(item.get("degree")),
                            _inline(item.get("tech_stack")),
                        ],
                    )
                )
                if meta:
                    doc.add_paragraph(meta)
                for field in ["description", "highlights", "content"]:
                    if item.get(field):
                        doc.add_paragraph(_plain(item.get(field)))
        elif key.startswith("custom_"):
            for item in section.get("items", []):
                doc.add_heading(item.get("title") or "内容", level=2)
                doc.add_paragraph(_plain(item.get("content")))

    file_name = f"vitaflow_resume_{resume.id}.docx"
    file_path = settings.export_path / file_name
    doc.save(file_path)
    db.add(ExportRecord(user_id=user_id, resume_id=resume.id, file_type="word", file_name=file_name, file_path=str(file_path)))
    db.commit()
    return file_path
